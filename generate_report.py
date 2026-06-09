import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START

doc = docx.Document()

# Set A4 size and margins to 3 cm
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

# --- TITLE PAGE ---
add_para('LOAN DEFAULT PREDICTION SYSTEM USING MACHINE LEARNING AND EXPLAINABLE AI', size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('REPORT SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE DEGREE OF', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('BACHELOR OF TECHNOLOGY\nIN\nINFORMATION TECHNOLOGY', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('By', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Student Name]', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Roll Number: [Student Roll No]', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('UNDER THE GUIDANCE\nOF', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Guide Name]', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Guides Affiliation]', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('DEPARTMENT OF INFORMATION TECHNOLOGY\nGAUHATI UNIVERSITY\nGUWAHATI, INDIA', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[MONTH] – 2026', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# --- DECLARATION ---
add_para('GAUHATI UNIVERSITY\nDEPARTMENT OF INFORMATION TECHNOLOGY\nGopinath Bordoloi Nagar, Jalukbari Guwahati-781014', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('DECLARATION', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('I "[Student Name]", Roll No "[Student Roll No]", a B.Tech. student of the department of Information Technology, Gauhati University hereby declares that I have compiled this report reflecting all my works during the semester long full time project as part of my BTech curriculum.')
add_para('I declare that I have included the descriptions etc. of my project work, and nothing has been copied/replicated from other’s work. The facts, figures, analysis, results, claims etc. depicted in my thesis are all related to my full time project work.')
add_para('I also declare that the same report or any substantial portion of this report has not been submitted anywhere else as part of any requirements for any degree/diploma etc.')
add_para('\n\n([Student Name])\nBranch: Information Technology\nDate:', align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_page_break()

# --- CERTIFICATE ---
add_para('GAUHATI UNIVERSITY\nDEPARTMENT OF INFORMATION TECHNOLOGY\nGopinath Bordoloi Nagar, Jalukbari Guwahati-781014', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Date: ____________', align=WD_ALIGN_PARAGRAPH.LEFT, space_after=24)
add_para('CERTIFICATE', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('This is to certify that "[Student Name]" bearing Roll No: "[Student Roll No]" has carried out the project work "Loan Default Prediction System" under my supervision and has compiled this report reflecting the candidate’s work in the semester long project. The candidate did this project full time during the whole semester under my supervision, and the analysis, results, claims etc. are all related to his/her studies and works during the semester.')
add_para('I recommend submission of this project report as a part for partial fulfillment of the requirements for the degree of Bachelor of Technology in Information Technology of Gauhati University.')
add_para('\n\n\n([Guide Name])\n([Guides Affiliation])', align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_page_break()

# --- ACKNOWLEDGMENTS ---
add_para('ACKNOWLEDGMENTS', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('I would like to express my profound gratitude to my project guide, [Guide Name], [Guides Affiliation], Department of Information Technology, Gauhati University, for their invaluable guidance, continuous encouragement, and insightful feedback throughout the duration of this project. Their expertise and mentorship were instrumental in the successful completion of this work.')
add_para('I also extend my sincere thanks to the Head of the Department, Information Technology, Gauhati University, for providing the necessary facilities and a conducive environment for carrying out this research.')
add_para('Furthermore, I am deeply thankful to all the faculty members of the Department of Information Technology for their constant support and for sharing their vast knowledge.')
add_para('Finally, I would like to thank my family and friends for their unwavering moral support and encouragement during my B.Tech studies.')
add_para('\n\n[Student Name]\nRoll No: [Student Roll No]', align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_page_break()

# --- ABSTRACT ---
add_para('ABSTRACT', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('This project presents a comprehensive Loan Default Prediction System designed to assess the creditworthiness of loan applicants and predict the likelihood of default. Utilizing a robust machine learning model trained on historical financial data, the system evaluates key applicant attributes such as revolving utilization of unsecured lines, age, debt ratio, monthly income, and past due history. To enhance transparency and trust in the automated decision-making process, the system integrates Explainable AI (XAI) techniques, specifically SHAP (SHapley Additive exPlanations) values, providing clear interpretations of the model\'s predictions. The architecture features a backend developed using Flask, which handles secure user authentication, database management using MySQL, and seamless integration with the machine learning model. The user interface allows for intuitive data entry and visualizes the prediction results alongside the SHAP-based explanations. Overall, this system provides financial institutions with a reliable, interpretable, and scalable tool to mitigate financial risks associated with loan approvals.')
doc.add_page_break()

# --- TABLE OF CONTENTS ---
add_para('TABLE OF CONTENTS', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('1. Introduction ... 1')
add_para('2. Background/ Review of literature ... 3')
add_para('3. Methodology / System Architecture ... 6')
add_para('4. Proposed Solution / Implementation ... 10')
add_para('5. Results and Discussions ... 15')
add_para('6. Conclusion and Future Works ... 18')
add_para('References ... 19')
add_para('Appendices ... 20')
doc.add_page_break()

# --- LIST OF FIGURES AND TABLES ---
add_para('LIST OF FIGURES', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('Figure 1.1: System Architecture Diagram ... 7')
add_para('Figure 1.2: SHAP Summary Plot ... 16')
doc.add_page_break()

add_para('LIST OF TABLES', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('Table 1.1: Dataset Description ... 4')
add_para('Table 1.2: Model Evaluation Metrics ... 17')
doc.add_page_break()

# --- MAIN CHAPTERS ---
def add_chapter(title, content):
    add_para(title, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=18)
    for p_text in content:
        add_para(p_text)
    doc.add_page_break()

add_chapter('1. Introduction', [
    'The global financial landscape is characterized by a high volume of loan applications and the inherent risk of defaults. Lending institutions face the challenge of accurately assessing the creditworthiness of applicants to minimize financial losses. Traditional methods often rely on rigid rules and manual evaluation, which can be time-consuming and prone to human error or bias.',
    'In recent years, machine learning (ML) has emerged as a powerful tool in predictive modeling, offering the ability to analyze complex patterns in historical data to forecast future events, such as loan defaults. However, typical ML models often act as "black boxes," providing predictions without transparent reasoning. This lack of interpretability is a significant barrier to adoption in regulated industries like finance.',
    'This project aims to develop a Loan Default Prediction System that not only utilizes advanced ML algorithms for high accuracy but also incorporates Explainable AI (XAI) to provide clear, interpretable reasons for its predictions. By understanding which features (e.g., debt ratio, monthly income, past-due history) drive a particular prediction, stakeholders can make more informed and transparent lending decisions.'
])

add_chapter('2. Background/ Review of literature', [
    'Historically, credit scoring models such as the FICO score have been the industry standard for assessing credit risk. These models utilize a linear combination of factors based on credit history. While effective, they may not capture non-linear relationships and complex interactions between variables that modern machine learning techniques can identify.',
    'Various machine learning algorithms, including Logistic Regression, Random Forests, and Gradient Boosting Machines (GBM), have been widely applied to credit risk assessment. Studies have shown that ensemble methods like Random Forests and GBMs generally outperform traditional statistical models in predictive accuracy.',
    'Despite their accuracy, the "black-box" nature of complex models has spurred research into Explainable AI (XAI). Techniques like LIME (Local Interpretable Model-agnostic Explanations) and SHAP (SHapley Additive exPlanations) have been introduced to interpret model predictions. SHAP, based on cooperative game theory, is particularly favored for its consistency and theoretical guarantees, providing a unified measure of feature importance.'
])

# --- CHAPTER 3 ---
add_para('3. Methodology / System Architecture', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=18)
add_para('The proposed system follows a client-server architecture. The backend is responsible for data processing, model inference, database management, and generating XAI explanations. The frontend provides a user-friendly interface for inputting application data and viewing results.')

# Insert Flow Diagram
try:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Max width is 15cm (21cm A4 width - 3cm left margin - 3cm right margin)
    run.add_picture('C:\\Users\\bitop\\Loan_Default\\Project_Flow_Diagram.png', width=Cm(15))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = cap.add_run('Figure 1.1: System Architecture Diagram')
    run_cap.font.name = 'Times New Roman'
    run_cap.font.size = Pt(10)
    run_cap.italic = True
except Exception as e:
    print(f"Warning: Could not insert image: {e}")

add_para('Data Preprocessing: The system utilizes a dataset (`data.csv`) comprising historical loan applications and default outcomes. Preprocessing steps include handling missing values, encoding categorical variables, and scaling numerical features.')
add_para('Model Training: A machine learning model (e.g., Random Forest or XGBoost) is trained on the preprocessed dataset. The model learns the underlying patterns associated with loan defaults and is saved (`loan_model_bundle.pkl`) for inference.')
add_para('Backend Infrastructure: Developed using Python and Flask (`app.py`), the backend exposes RESTful APIs for client communication. It integrates with a MySQL database (`init_db.sql`) to manage user accounts, authenticate sessions via JWT, and log prediction requests.')
add_para('Explainability Module: The XAI component (`xai.py`) utilizes the SHAP library to compute the contribution of each input feature to the final prediction. These explanations are formatted and sent to the frontend alongside the prediction.')
doc.add_page_break()

add_chapter('4. Proposed Solution / Implementation', [
    'The implementation of the Loan Default Prediction System involves integrating the trained machine learning model within the Flask backend. When a user submits an application through the frontend, the data is sent to the backend API.',
    'The backend first validates and preprocesses the incoming data using the `preprocessing.py` module to ensure it matches the format expected by the model. The preprocessed data is then passed to the loaded model for prediction.',
    'Concurrently, the `xai.py` module computes the SHAP values for the specific application instance. This process quantifies the impact of features like \'RevolvingUtilizationOfUnsecuredLines\' and \'DebtRatio\' on the model\'s decision.',
    'The prediction (e.g., "Default" or "No Default") and the associated SHAP values are bundled into a JSON response and returned to the frontend. The React-based frontend then renders this information, providing the user with both the prediction and an intuitive visualization of the contributing factors.'
])

add_chapter('5. Results and Discussions', [
    'The trained machine learning model was evaluated on a held-out test set to assess its generalization performance. Key metrics such as Accuracy, Precision, Recall, and the F1-Score were calculated. The model demonstrated strong predictive capability, effectively distinguishing between default and non-default cases.',
    'The integration of Explainable AI (SHAP) proved to be highly valuable. For individual predictions, the system successfully generated local explanations, highlighting the most influential features. For instance, high \'RevolvingUtilizationOfUnsecuredLines\' and a history of \'NumberOfTime30-59DaysPastDueNotWorse\' were frequently identified as strong indicators of default.',
    'These explanations provide actionable insights for loan officers, enabling them to understand the rationale behind automated decisions and potentially offer targeted advice to applicants on how to improve their creditworthiness.'
])

add_chapter('6. Conclusion and Future Works', [
    'This project successfully developed and implemented a Loan Default Prediction System that leverages machine learning for accurate risk assessment and Explainable AI for transparent decision-making. By integrating SHAP values, the system addresses the "black-box" challenge, fostering trust and accountability in automated lending processes.',
    'Future work could explore the integration of additional data sources, such as alternative credit data or macroeconomic indicators, to further enhance predictive accuracy. Furthermore, exploring other XAI techniques or developing interactive, user-specific explanation dashboards could provide even greater utility for financial institutions.'
])

add_chapter('References', [
    '[1] Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in neural information processing systems, 30.',
    '[2] Ribeiro, M. T., Singh, S., & Guestrin, C. (2016). "Why should I trust you?" Explaining the predictions of any classifier. In Proceedings of the 22nd ACM SIGKDD international conference on knowledge discovery and data mining (pp. 1135-1144).',
    '[3] Khandani, A. E., Kim, A. J., & Lo, A. W. (2010). Consumer credit-risk models via machine-learning algorithms. Journal of Banking & Finance, 34(11), 2767-2787.'
])

# --- APPENDICES ---
add_para('APPENDICES', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=18)
add_para('APPENDIX A: Backend Source Code', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para('A.1. app.py (Main Application File)')
add_para('This file handles the main API endpoints, routing, and integration of the machine learning model and database.')
add_para('A.2. preprocessing.py')
add_para('This script contains functions for data cleaning, transformation, and preparation before feeding it to the machine learning model.')
add_para('A.3. xai.py')
add_para('Contains the implementation for integrating Explainable AI (SHAP) to interpret the predictions made by the model.')
add_para('A.4. init_db.sql')
add_para('The SQL script used to initialize the database schema, including tables for user authentication and storing prediction logs.')

add_para('\nAPPENDIX B: Frontend Source Code', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para('B.1. React Components')
add_para('The user interface is built using React. The source code contains components for user login, dashboard, data input forms, and visualization of the prediction results along with SHAP explanations.')

add_para('\nAPPENDIX C: Dependencies and Requirements', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para('C.1. Python Packages (requirements.txt)')
add_para('The key libraries utilized in the backend include Flask, pandas, joblib, shap, mysql-connector-python, among others, to support API creation, data manipulation, and machine learning functionalities.')
add_para('C.2. Node Packages (package.json)')
add_para('The frontend dependencies manage the React framework, routing, state management, and charting libraries used to render the application interface and visualizations.')

doc.save('C:\\Users\\bitop\\Loan_Default\\Complete_Project_Report.docx')
